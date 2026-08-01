import csv, json
from pathlib import Path
from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook

ALLOWED = {".pdf", ".docx", ".xlsx", ".xlsm", ".txt", ".csv", ".json", ".md"}

def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return "\n\n".join((p.extract_text() or "") for p in PdfReader(str(path)).pages).strip()
    if ext == ".docx":
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            parts.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(parts).strip()
    if ext in {".xlsx", ".xlsm"}:
        wb = load_workbook(str(path), read_only=True, data_only=True)
        out = []
        for ws in wb.worksheets:
            out.append(f"# Sheet: {ws.title}")
            for row in ws.iter_rows(values_only=True):
                out.append(" | ".join("" if v is None else str(v) for v in row))
        return "\n".join(out).strip()
    if ext == ".csv":
        with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as f:
            return "\n".join(" | ".join(row) for row in csv.reader(f))
    if ext == ".json":
        with path.open("r", encoding="utf-8", errors="replace") as f:
            return json.dumps(json.load(f), ensure_ascii=False, indent=2)
    if ext in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Unsupported file type: {ext}")

def chunks(text: str, size: int = 1800, overlap: int = 250):
    text = " ".join(text.split())
    if not text:
        return []
    out, start = [], 0
    while start < len(text):
        end = min(len(text), start + size)
        if end < len(text):
            split = text.rfind(" ", start, end)
            if split > start + size // 2:
                end = split
        out.append(text[start:end].strip())
        if end >= len(text): break
        start = max(start + 1, end - overlap)
    return out
